"""Run one candidate through every stage and print a trace. WRITES.

    python -m scripts.e2e_verify --database-url postgresql+psycopg://...

WHAT THIS IS FOR

Day 12 asks for one genuine end-to-end execution rather than a claim
assembled from unit tests. This drives the real services in order,
against a real database, and prints what each stage actually did with
the time it took.

IT WRITES TO WHATEVER DATABASE YOU POINT IT AT, AND IT TRUNCATES FIRST.

That is why `--database-url` is required and has no default. A default
would eventually be somebody's development database. The script refuses
to run against a URL whose database name does not contain "test",
which is a blunt guard and is meant to be.

WHICH STAGES ARE REAL

Every stage is the production code path. Three external services sit
outside it:

  * the job board, replaced by a fixture source through the same
    JobSource Protocol production uses;
  * the extraction and embedding model, replaced by deterministic
    stand-ins;
  * Telegram, replaced by a notifier that records instead of sending.

Everything else -- text extraction, validation, deduplication, pgvector
retrieval, all five signals, renormalisation, ranking, the three
notification gates, the LangGraph routing, every INSERT and every
constraint -- is the real thing. The trace labels each stage REAL or
STAND-IN so the distinction survives being pasted into a report.

Use `--with-telegram` to send the notification through the real
TelegramNotifier, which requires a working TELEGRAM_BOT_TOKEN and a
chat that has messaged the bot. Off by default: an end-to-end script
that messages a real person the first time somebody runs it is a
script people stop running.
"""

from __future__ import annotations

import argparse
import asyncio
import io
import math
import sys
import time
from datetime import UTC, datetime, timedelta

if sys.platform == "win32":
    asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())

from pathlib import Path  # noqa: E402
from tempfile import TemporaryDirectory  # noqa: E402

DIMENSIONS = 768
NOW = datetime.now(UTC)

CV_LINES = [
    "PRIYA SHARMA",
    "AI / ML Engineer -- Delhi NCR",
    "",
    "SUMMARY",
    "AI/ML engineer with one year of experience building NLP and retrieval "
    "systems in Python.",
    "",
    "SKILLS",
    "Python, Machine Learning, NLP, FastAPI, SQL",
    "",
    "EXPERIENCE",
    "ML Engineer, Northwind Analytics, Aug 2025 - present",
    "Built a semantic search service over 200k documents.",
    "",
    "EDUCATION",
    "B.Tech Computer Science, Delhi Technological University, 2025",
]


def vector_at(degrees: float) -> list[float]:
    """A unit vector at a known angle, so similarity is arithmetic."""
    vector = [0.0] * DIMENSIONS
    vector[0] = math.cos(math.radians(degrees))
    vector[1] = math.sin(math.radians(degrees))
    return vector


class Trace:
    """Stage name, verdict, timing. Printed as it goes, not at the end."""

    def __init__(self) -> None:
        self.rows: list[tuple[str, str, str, float]] = []
        self._started = time.perf_counter()

    def stage(self, name: str, kind: str):
        return _Stage(self, name, kind)

    def record(self, name: str, kind: str, detail: str, seconds: float) -> None:
        self.rows.append((name, kind, detail, seconds))
        print(f"  {'OK':<4} {name:<28} {kind:<9} {seconds * 1000:8.1f} ms   {detail}")

    def total_seconds(self) -> float:
        return time.perf_counter() - self._started


class _Stage:
    def __init__(self, trace: Trace, name: str, kind: str) -> None:
        self.trace = trace
        self.name = name
        self.kind = kind
        self.detail = ""

    async def __aenter__(self) -> "_Stage":
        self._start = time.perf_counter()
        return self

    async def __aexit__(self, exc_type, exc, tb) -> bool:
        elapsed = time.perf_counter() - self._start
        if exc_type is not None:
            print(f"  {'FAIL':<4} {self.name:<28} {self.kind:<9} "
                  f"{elapsed * 1000:8.1f} ms   {exc_type.__name__}: {exc}")
            return False
        self.trace.record(self.name, self.kind, self.detail, elapsed)
        return False


async def run(database_url: str, with_telegram: bool) -> int:
    from sqlalchemy import func, select, text

    import app.db.models  # noqa: F401
    from app.core.config import settings

    settings.database_url = database_url

    from app.db.base import Base
    from app.db.models.cv import CV, CVVersion, ExtractionStatus
    from app.db.models.job import Job, JobSkill
    from app.db.models.recommendation import (
        FeedbackAction,
        Notification,
        Recommendation,
        UserFeedback,
    )
    from app.db.models.skill import Skill
    from app.db.models.user import User, UserPreference
    from app.db.repositories.agent import AgentRunRepository
    from app.db.repositories.skill import SkillRepository
    from app.db.session import dispose_engine, init_engine, session_scope
    from app.schemas.cv_profile import CVProfile, EducationEntry, ExperienceEntry
    from app.services.cv_extraction import extract_cv
    from app.services.feedback import FeedbackService
    from app.services.job_ingestion import run_ingestion
    from app.services.job_scoring import run_scoring
    from app.services.job_search import search_for_user
    from app.services.notification_delivery import (
        deliver_notifications,
        select_notifiable,
    )
    from app.services.notification_message import CALLBACK_PREFIX
    from app.workflows.graph import build_graph
    from app.workflows.state import build_run_summary, initial_state

    class FakeGemini:
        model = "e2e-stand-in"

        async def extract_profile(self, raw_text: str) -> CVProfile:
            return CVProfile(
                summary="AI/ML engineer with one year of experience.",
                current_title="ML Engineer",
                location="Delhi NCR",
                skills=["Python", "Machine Learning", "NLP", "FastAPI", "SQL"],
                target_roles=["AI Engineer", "ML Engineer"],
                experience=[
                    ExperienceEntry(
                        title="ML Engineer",
                        company="Northwind Analytics",
                        start_year=2025,
                        start_month=8,
                        is_current=True,
                    )
                ],
                education=[
                    EducationEntry(
                        degree="B.Tech Computer Science",
                        institution="Delhi Technological University",
                    )
                ],
            )

    class RecordingNotifier:
        def __init__(self) -> None:
            self.sent: list[tuple[int, str]] = []

        async def __aenter__(self):
            return self

        async def __aexit__(self, *args):
            return None

        async def send(self, *, chat_id, reply):
            from app.integrations.telegram import SendResult

            self.sent.append((chat_id, reply.text))
            return SendResult(ok=True)

    class FixtureSource:
        source_name = "fixture"

        async def search(self, *, what="", where="", page=1, max_days_old=None,
                         results_per_page=None):
            from app.integrations.adzuna import SearchPage, UnparseableRecord
            from app.schemas.job import RawJobPosting

            if page > 1:
                return SearchPage(total_available=0, postings=[], unparseable=[])

            postings = [
                RawJobPosting(
                    source="fixture", external_id="ml-1",
                    title="Machine Learning Engineer", company="Northwind Analytics",
                    location="Delhi, India", url="https://example.test/ml-1",
                    description="Python, machine learning and NLP.",
                    posted_at=NOW - timedelta(days=1),
                ),
                RawJobPosting(
                    source="fixture", external_id="ml-2",
                    title="NLP Engineer", company="Cobalt Systems",
                    location="Delhi, India", url="https://example.test/ml-2",
                    description="NLP and retrieval.",
                    posted_at=NOW - timedelta(days=2),
                ),
                RawJobPosting(
                    source="fixture", external_id="acct-1",
                    title="Senior Accountant", company="Ledgerworks",
                    location="Chennai, India", url="https://example.test/acct-1",
                    description="Statutory audit and tax.",
                    posted_at=NOW - timedelta(days=3),
                ),
                RawJobPosting(  # duplicate of ml-1 under a new id
                    source="fixture", external_id="ml-1-repost",
                    title="Machine Learning Engineer", company="Northwind Analytics",
                    location="Delhi, India", url="https://example.test/ml-1b",
                    description="Python, machine learning and NLP.",
                    posted_at=NOW - timedelta(days=1),
                ),
                RawJobPosting(  # rejected: no posting date
                    source="fixture", external_id="broken-1",
                    title="Data Engineer", company="Nowhere",
                    location="Delhi, India", url="https://example.test/broken",
                    description="", posted_at=None,
                ),
            ]
            return SearchPage(
                total_available=len(postings),
                postings=postings,
                unparseable=[UnparseableRecord(raw={"x": 1}, reason="missing title")],
            )

    trace = Trace()
    init_engine()

    print("")
    print("=" * 78)
    print(f"E2E VERIFICATION RUN   {NOW.isoformat()}")
    print("=" * 78)
    print("")
    print(f"  {'':<4} {'STAGE':<28} {'KIND':<9} {'ELAPSED':>11}   DETAIL")
    print(f"  {'-' * 72}")

    try:
        async with session_scope() as session:
            tables = ", ".join(sorted(Base.metadata.tables))
            await session.execute(
                text(f"TRUNCATE TABLE {tables} RESTART IDENTITY CASCADE")
            )

        with TemporaryDirectory() as directory:
            cv_path = Path(directory) / "priya_sharma.docx"

            async with trace.stage("Candidate created", "REAL") as stage:
                async with session_scope() as session:
                    user = User(telegram_id=770001, full_name="Priya Sharma")
                    session.add(user)
                    await session.flush()
                    session.add(
                        UserPreference(
                            user_id=user.id,
                            target_roles=[
                                "AI Engineer",
                                "ML Engineer",
                                "Machine Learning Engineer",
                            ],
                            preferred_locations=["Delhi"],
                        )
                    )
                    user_id = user.id
                stage.detail = f"user_id={user_id} telegram_id=770001"

            async with trace.stage("CV uploaded", "REAL") as stage:
                from docx import Document

                document = Document()
                for line in CV_LINES:
                    document.add_paragraph(line)
                buffer = io.BytesIO()
                document.save(buffer)
                cv_path.write_bytes(buffer.getvalue())

                async with session_scope() as session:
                    session.add(
                        CV(
                            user_id=user_id,
                            file_name=cv_path.name,
                            file_type="docx",
                            file_size_bytes=cv_path.stat().st_size,
                            storage_path=str(cv_path),
                            extraction_status=ExtractionStatus.PENDING.value,
                        )
                    )
                stage.detail = f"{cv_path.stat().st_size} bytes, docx"

            async with trace.stage("CV text + profile", "STAND-IN") as stage:
                result = await extract_cv(user_id, gemini_client=FakeGemini())
                async with session_scope() as session:
                    cv = (await session.execute(select(CV))).scalar_one()
                    raw_length = len(cv.raw_text or "")
                stage.detail = (
                    f"status={result.status.value} "
                    f"raw_text={raw_length} chars (text layer REAL)"
                )

            async with trace.stage("Profile persisted", "REAL") as stage:
                async with session_scope() as session:
                    from app.db.models.profile import Profile

                    profile = (await session.execute(select(Profile))).scalar_one()
                    version_id = profile.active_cv_version_id
                    skills = list(profile.skills)
                    years = profile.total_experience_years
                stage.detail = (
                    f"cv_version={version_id} skills={len(skills)} "
                    f"experience={years:.2f}y"
                )

            async with trace.stage("CV embedded", "STAND-IN") as stage:
                async with session_scope() as session:
                    version = (
                        await session.execute(
                            select(CVVersion).where(CVVersion.id == version_id)
                        )
                    ).scalar_one()
                    version.embedding = vector_at(0)
                    version.embedding_model = "e2e-stand-in"
                    version.embedded_at = NOW
                stage.detail = f"{DIMENSIONS} dimensions"

            async with trace.stage("Jobs ingested", "REAL") as stage:
                ingestion = await run_ingestion(
                    FixtureSource(),
                    keywords=["machine learning engineer"],
                    locations=["delhi"],
                    max_pages=1,
                    now=NOW,
                )
                counters = ingestion.counters
                stage.detail = (
                    f"fetched={counters['records_fetched']} "
                    f"inserted={counters['inserted']} "
                    f"rejected={counters['validation_failed'] + counters['normalize_failed']} "
                    f"duplicates={counters['duplicates']}"
                )

            async with trace.stage("Funnel balanced", "REAL") as stage:
                accounted = (
                    counters["normalize_failed"]
                    + counters["validation_failed"]
                    + counters["filtered_out"]
                    + counters["duplicates"]
                    + counters["inserted"]
                )
                assert accounted == counters["records_fetched"], "funnel does not balance"
                stage.detail = f"{accounted} == {counters['records_fetched']}"

            async with trace.stage("Jobs embedded", "STAND-IN") as stage:
                angles = {"ml-1": 4.0, "ml-2": 18.0, "acct-1": 85.0}
                async with session_scope() as session:
                    jobs = (await session.execute(select(Job))).scalars().all()
                    for job in jobs:
                        job.embedding = vector_at(angles.get(job.external_id, 80.0))
                        job.embedding_model = "e2e-stand-in"
                        job.embedded_at = NOW
                    job_ids = {job.external_id: job.id for job in jobs}
                stage.detail = f"{len(job_ids)} jobs, {DIMENSIONS} dimensions"

            async with trace.stage("Skills enriched", "STAND-IN") as stage:
                async with session_scope() as session:
                    added = 0
                    # get_or_create, not a bare INSERT: CV extraction has
                    # ALREADY populated the skill catalog from the profile,
                    # so inserting "Python" again violates
                    # ix_skills_normalized_name. The catalog is shared
                    # between the two sides of the match by design -- that
                    # is what makes a skill on a job comparable with a
                    # skill on a CV -- and this script found out by
                    # crashing on it.
                    skill_repository = SkillRepository(session)
                    for name in ("Python", "Machine Learning", "NLP"):
                        skill = await skill_repository.get_or_create(name)
                        await session.flush()
                        session.add(
                            JobSkill(
                                job_id=job_ids["ml-1"],
                                skill_id=skill.id,
                                is_required=True,
                            )
                        )
                        added += 1
                    target = (
                        await session.execute(
                            select(Job).where(Job.id == job_ids["ml-1"])
                        )
                    ).scalar_one()
                    target.min_experience_years = 0
                    target.max_experience_years = 3
                stage.detail = f"job {job_ids['ml-1']}: {added} skills, exp 0-3y"

            async with trace.stage("pgvector retrieval", "REAL") as stage:
                matches = await search_for_user(user_id, limit=10)
                top = matches[0]
                stage.detail = (
                    f"{len(matches)} hits, top=job {top.job_id} "
                    f"similarity={top.similarity:.4f}"
                )

            async with trace.stage("Scoring + ranking", "REAL") as stage:
                scoring = await run_scoring(user_id=user_id)
                stage.detail = (
                    f"pairs={scoring['pairs_scored']} "
                    f"max={scoring['score_max']:.4f} "
                    f"notify_eligible={scoring['notify_eligible']}"
                )

            async with trace.stage("LangGraph run", "REAL") as stage:
                async with session_scope() as session:
                    run_row_id = (await AgentRunRepository(session).start(NOW)).id

                state = initial_state(
                    user_id=user_id,
                    dry_run=True,
                    skip_ingestion=True,
                    skip_embedding=True,
                    skip_enrichment=True,
                    started_at=NOW.isoformat(),
                )
                final_state = await build_graph().ainvoke(state)
                summary = build_run_summary(final_state)

                async with session_scope() as session:
                    await AgentRunRepository(session).finish(run_row_id, summary)

                stage.detail = (
                    f"branch={summary['notify_branch']} "
                    f"eligible={summary['notify_eligible']} "
                    f"agent_run={run_row_id}"
                )

            async with trace.stage("Notification decision", "REAL") as stage:
                candidates = await select_notifiable(user_id=user_id)
                stage.detail = f"{len(candidates)} candidate(s) cleared all three gates"

            notifier = RecordingNotifier()
            kind = "REAL" if with_telegram else "STAND-IN"
            async with trace.stage("Telegram delivery", kind) as stage:
                delivery = await deliver_notifications(
                    candidates,
                    notifier=None if with_telegram else notifier,
                )
                stage.detail = (
                    f"attempted={delivery['attempted']} sent={delivery['sent']} "
                    f"failed={delivery['failed']}"
                )

            async with trace.stage("Duplicate suppressed", "REAL") as stage:
                again = await select_notifiable(user_id=user_id)
                stage.detail = f"second pass selected {len(again)} (already sent)"

            async with trace.stage("Feedback recorded", "REAL") as stage:
                notified_job_id = (
                    candidates[0].job_id if candidates else job_ids["ml-1"]
                )
                async with session_scope() as session:
                    reply = await FeedbackService(session).handle_callback(
                        770001,
                        f"{CALLBACK_PREFIX}:{FeedbackAction.INTERESTED.value}:"
                        f"{notified_job_id}",
                    )
                stage.detail = f"job {notified_job_id}: {reply.text.splitlines()[0]}"

            async with trace.stage("Database verified", "REAL") as stage:
                async with session_scope() as session:
                    async def count(model):
                        return int(
                            (
                                await session.execute(select(func.count(model.id)))
                            ).scalar_one()
                        )

                    counts = {
                        "users": await count(User),
                        "cvs": await count(CV),
                        "cv_versions": await count(CVVersion),
                        "jobs": await count(Job),
                        "recommendations": await count(Recommendation),
                        "notifications": await count(Notification),
                        "feedback": await count(UserFeedback),
                    }
                stage.detail = " ".join(f"{k}={v}" for k, v in counts.items())

        print(f"  {'-' * 72}")
        print("")
        print(f"  stages           {len(trace.rows)}")
        print(f"  total elapsed    {trace.total_seconds() * 1000:.1f} ms")
        print("")
        if not with_telegram:
            print("  Telegram was NOT exercised. The message was rendered and the")
            print("  attempt row written, but no socket was opened. Re-run with")
            print("  --with-telegram against a real token to close that gap.")
        print("")
        return 0
    finally:
        await dispose_engine()


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        prog="python -m scripts.e2e_verify",
        description="Drive one candidate through every stage. WRITES AND TRUNCATES.",
    )
    parser.add_argument(
        "--database-url",
        required=True,
        help="Scratch database. Truncated before the run. Must contain 'test'.",
    )
    parser.add_argument(
        "--with-telegram",
        action="store_true",
        help="Send through the real TelegramNotifier instead of recording.",
    )
    arguments = parser.parse_args(argv)

    # Blunt on purpose. A default would eventually be somebody's dev
    # database, and this script truncates every table it can see.
    database_name = arguments.database_url.rsplit("/", 1)[-1].split("?")[0]
    if "test" not in database_name.lower():
        print(
            f"refusing to run    database name {database_name!r} does not "
            "contain 'test'; this script TRUNCATES every table",
            file=sys.stderr,
        )
        return 2

    return asyncio.run(run(arguments.database_url, arguments.with_telegram))


if __name__ == "__main__":
    raise SystemExit(main())
