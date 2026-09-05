"""Day 12 integration test 1 -- CV -> extraction -> profile -> database.

    CV file on disk
        -> app.services.cv_text.extract_raw_text   (REAL, python-docx)
        -> app.services.cv_extraction.extract_cv   (REAL)
        -> cv_versions + profiles + skills rows     (REAL PostgreSQL)

The .docx is generated here and read back through the real extractor,
so the text layer is genuinely parsed rather than injected. The only
substitution is `GeminiClient`, which is a paid network call: the fake
returns a `CVProfile` that a real extraction would plausibly produce.
That boundary is stated in docs/TEST_RESULTS.md rather than implied,
because "CV extraction verified" and "CV extraction verified except
for the model" are different claims.
"""

from __future__ import annotations

import io
from collections.abc import Awaitable, Callable
from pathlib import Path
from typing import Any

from docx import Document
from sqlalchemy import select

from app.db.models.cv import CV, CVVersion, ExtractionStatus
from app.db.models.profile import Profile
from app.db.models.skill import Skill
from app.db.models.user import User
from app.db.session import session_scope
from app.schemas.cv_profile import CVProfile, EducationEntry, ExperienceEntry
from app.services.cv_extraction import extract_cv

CV_TEXT = [
    "PRIYA SHARMA",
    "AI / ML Engineer -- Delhi NCR",
    "",
    "SUMMARY",
    "AI/ML engineer with one year of experience building NLP and "
    "retrieval systems in Python.",
    "",
    "SKILLS",
    "Python, Machine Learning, NLP, FastAPI, SQL, PostgreSQL",
    "",
    "EXPERIENCE",
    "ML Engineer, Northwind Analytics, Aug 2025 - present",
    "Built a semantic search service over 200k documents.",
    "",
    "EDUCATION",
    "B.Tech Computer Science, Delhi Technological University, 2025",
]


class FakeGeminiClient:
    """Stands in for the paid extraction call. Nothing else is faked.

    Shaped as an ordinary object rather than a subclass, matching the
    reason `extract_cv` accepts `gemini_client` as a parameter at all.
    """

    model = "fake-extraction-model"

    def __init__(self) -> None:
        self.seen_text: str | None = None

    async def extract_profile(self, raw_text: str) -> CVProfile:
        self.seen_text = raw_text
        return CVProfile(
            summary=(
                "AI/ML engineer with one year of experience building NLP "
                "and retrieval systems in Python."
            ),
            current_title="ML Engineer",
            location="Delhi NCR",
            skills=["Python", "Machine Learning", "NLP", "FastAPI", "SQL"],
            target_roles=["AI Engineer", "ML Engineer"],
            experience=[
                ExperienceEntry(
                    title="ML Engineer",
                    company="Northwind Analytics",
                    start_year=2025,
                    start_month=8,
                    is_current=True,
                )
            ],
            education=[
                EducationEntry(
                    degree="B.Tech Computer Science",
                    institution="Delhi Technological University",
                )
            ],
        )


def _write_docx(path: Path) -> None:
    document = Document()
    for line in CV_TEXT:
        document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)
    path.write_bytes(buffer.getvalue())


def test_a_cv_becomes_a_persisted_profile(
    run_with_database: Callable[[Callable[[], Awaitable[Any]]], Any],
    tmp_path: Path,
) -> None:
    storage_path = tmp_path / "priya_cv.docx"
    _write_docx(storage_path)

    client = FakeGeminiClient()

    async def body() -> dict[str, Any]:
        async with session_scope() as session:
            user = User(telegram_id=900001, full_name="Priya Sharma")
            session.add(user)
            await session.flush()
            session.add(
                CV(
                    user_id=user.id,
                    file_name="priya_cv.docx",
                    file_type="docx",
                    file_size_bytes=storage_path.stat().st_size,
                    storage_path=str(storage_path),
                    extraction_status=ExtractionStatus.PENDING.value,
                )
            )
            user_id = user.id

        result = await extract_cv(user_id, gemini_client=client)

        async with session_scope() as session:
            cv = (await session.execute(select(CV))).scalar_one()
            version = (await session.execute(select(CVVersion))).scalar_one()
            profile = (await session.execute(select(Profile))).scalar_one()
            skills = (
                (await session.execute(select(Skill.normalized_name))).scalars().all()
            )
            return {
                "result": result,
                "cv_status": cv.extraction_status,
                "cv_raw_text": cv.raw_text,
                "cv_extracted_at": cv.extracted_at,
                "version_id": version.id,
                "version_cv_id": version.cv_id,
                "version_number": version.version,
                "extracted_profile": version.extracted_profile,
                "extraction_model": version.extraction_model,
                "profile_user_id": profile.user_id,
                "profile_summary": profile.summary,
                "profile_title": profile.current_title,
                "profile_years": profile.total_experience_years,
                "profile_skills": profile.skills,
                "active_cv_version_id": profile.active_cv_version_id,
                "catalog_skills": sorted(skills),
                "user_id": user_id,
            }

    observed = run_with_database(body)

    # The real text extractor ran: the model saw the words from the file.
    assert client.seen_text is not None
    assert "Northwind Analytics" in client.seen_text
    assert "Delhi Technological University" in client.seen_text

    assert observed["result"].status is ExtractionStatus.COMPLETE
    assert observed["cv_status"] == ExtractionStatus.COMPLETE.value
    assert observed["cv_extracted_at"] is not None
    assert "PRIYA SHARMA" in observed["cv_raw_text"]

    # Version row, and the profile pointing at it. These two being
    # consistent is the whole reason the embedding can never go stale.
    assert observed["version_number"] == 1
    assert observed["extraction_model"] == "fake-extraction-model"
    assert observed["active_cv_version_id"] == observed["version_id"]

    assert observed["profile_user_id"] == observed["user_id"]
    assert observed["profile_title"] == "ML Engineer"

    # Skills on the profile are NORMALIZED catalog keys, while the
    # version keeps the model's spellings. Asserting both directions,
    # because the split is what stops a job asking for "nodejs" from
    # missing a CV that wrote "Node.js".
    assert "machine learning" in [s.lower() for s in observed["profile_skills"]]
    assert "Machine Learning" in observed["extracted_profile"]["skills"]
    assert "python" in observed["catalog_skills"]

    # Experience is computed from the structured dates, not taken from
    # the model. One month of 2025-08 to now is more than zero years.
    assert observed["profile_years"] is not None
    assert observed["profile_years"] > 0


def test_a_cv_with_no_text_layer_is_not_a_failure(
    run_with_database: Callable[[Callable[[], Awaitable[Any]]], Any],
    tmp_path: Path,
) -> None:
    """A scanned CV must reach NO_TEXT_LAYER, not FAILED.

    Day 12 section 27 asks for the empty-CV case. The distinction
    matters to the user: FAILED reads as "we broke", NO_TEXT_LAYER
    reads as "send a text PDF", and only one of those is actionable.
    """
    storage_path = tmp_path / "scan.docx"
    document = Document()  # no paragraphs at all
    buffer = io.BytesIO()
    document.save(buffer)
    storage_path.write_bytes(buffer.getvalue())

    client = FakeGeminiClient()

    async def body() -> dict[str, Any]:
        async with session_scope() as session:
            user = User(telegram_id=900002, full_name="Scanned CV")
            session.add(user)
            await session.flush()
            session.add(
                CV(
                    user_id=user.id,
                    file_name="scan.docx",
                    file_type="docx",
                    storage_path=str(storage_path),
                    extraction_status=ExtractionStatus.PENDING.value,
                )
            )
            user_id = user.id

        result = await extract_cv(user_id, gemini_client=client)

        async with session_scope() as session:
            cv = (await session.execute(select(CV))).scalar_one()
            versions = (await session.execute(select(CVVersion))).scalars().all()
            profiles = (await session.execute(select(Profile))).scalars().all()
            return {
                "result": result,
                "cv_status": cv.extraction_status,
                "versions": len(versions),
                "profiles": len(profiles),
            }

    observed = run_with_database(body)

    assert observed["result"].status is ExtractionStatus.NO_TEXT_LAYER
    assert observed["cv_status"] == ExtractionStatus.NO_TEXT_LAYER.value
    # No model call was made, so no version and no profile were written.
    assert client.seen_text is None
    assert observed["versions"] == 0
    assert observed["profiles"] == 0


def test_a_corrupted_file_fails_without_taking_the_process_with_it(
    run_with_database: Callable[[Callable[[], Awaitable[Any]]], Any],
    tmp_path: Path,
) -> None:
    """Day 12 section 14: a corrupted CV must degrade, not crash."""
    storage_path = tmp_path / "corrupt.docx"
    storage_path.write_bytes(b"this is not a docx file at all")

    client = FakeGeminiClient()

    async def body() -> dict[str, Any]:
        async with session_scope() as session:
            user = User(telegram_id=900003, full_name="Corrupt CV")
            session.add(user)
            await session.flush()
            session.add(
                CV(
                    user_id=user.id,
                    file_name="corrupt.docx",
                    file_type="docx",
                    storage_path=str(storage_path),
                    extraction_status=ExtractionStatus.PENDING.value,
                )
            )
            user_id = user.id

        result = await extract_cv(user_id, gemini_client=client)

        async with session_scope() as session:
            cv = (await session.execute(select(CV))).scalar_one()
            return {
                "result": result,
                "cv_status": cv.extraction_status,
                "cv_error": cv.extraction_error,
            }

    observed = run_with_database(body)

    assert observed["result"].status is ExtractionStatus.FAILED
    assert observed["cv_status"] == ExtractionStatus.FAILED.value
    # The reason is stored, not just logged. A status with no reason
    # sends the next person to the logs for a run that may be days old.
    assert observed["cv_error"]
    assert client.seen_text is None


def test_a_nul_byte_in_the_text_layer_reaches_the_database_cleaned(
    run_with_database: Callable[[Callable[[], Awaitable[Any]]], Any],
    tmp_path: Path,
    monkeypatch,
) -> None:
    """The production failure of 2026-09-05, executed against PostgreSQL.

    A real CV's PDF text layer contained `\\x00` where the "+" of an
    international phone number belonged. `extract_cv` died in its THIRD
    phase -- after the Gemini call had been spent -- with

        psycopg.DataError: PostgreSQL text fields cannot contain NUL

    Only a real database raises that. No unit test could have caught
    it, which is the same reason this directory exists.

    The pdf extractor is stubbed because a .docx cannot carry a NUL and
    building a PDF whose font tables produce one is not worth it; what
    matters here is that the value which reaches `cvs.raw_text` and
    `cv_versions.extracted_profile` is accepted by the columns.
    """
    import app.services.cv_text as cv_text

    storage_path = tmp_path / "varun_cv.pdf"
    storage_path.write_bytes(b"%PDF-1.4 stub")

    dirty = "VARUN NARAD\nCybersecurity Engineer\n\x00918700430994"
    monkeypatch.setattr(cv_text, "_extract_pdf_text", lambda data: dirty)

    client = FakeGeminiClient()

    async def body() -> dict[str, Any]:
        async with session_scope() as session:
            user = User(telegram_id=900004, full_name="Varun Narad")
            session.add(user)
            await session.flush()
            session.add(
                CV(
                    user_id=user.id,
                    file_name="varun_cv.pdf",
                    file_type="pdf",
                    storage_path=str(storage_path),
                    extraction_status=ExtractionStatus.PENDING.value,
                )
            )
            user_id = user.id

        result = await extract_cv(user_id, gemini_client=client)

        async with session_scope() as session:
            cv = (await session.execute(select(CV))).scalar_one()
            version = (await session.execute(select(CVVersion))).scalar_one()
            return {
                "result": result,
                "cv_status": cv.extraction_status,
                "raw_text": cv.raw_text,
                "extracted_profile": version.extracted_profile,
            }

    observed = run_with_database(body)

    assert observed["result"].status is ExtractionStatus.COMPLETE
    assert observed["cv_status"] == ExtractionStatus.COMPLETE.value

    # Stored, and clean. Before the fix this row did not exist at all.
    assert "\x00" not in observed["raw_text"]
    assert "918700430994" in observed["raw_text"]

    # The model saw the cleaned text, so the JSONB column is safe too.
    assert client.seen_text is not None
    assert "\x00" not in client.seen_text
    assert "\x00" not in str(observed["extracted_profile"])
