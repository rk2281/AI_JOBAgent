"""Tests for the Day 4 CV extraction pieces that don't need a database.

extract_cv itself (app.services.cv_extraction) is not tested here. It
opens a real session and writes CV, CVVersion, Profile, and Skill
rows — testing it properly means testing against Postgres, which this
project's models already require (JSONB, pgvector's Vector type,
HNSW indexes have no SQLite equivalent). That verification happens
live, via scripts/extract_cv.py against the real database, the same
way onboarding's restart resilience and edge cases are verified via
scripts/*_dryrun.py rather than pytest. What's tested here is
everything underneath extract_cv that has no such dependency: the
schema, the skill-name normalization, and the file-format text
extraction.
"""

from __future__ import annotations

import pytest

from app.db.repositories.skill import normalize_skill_name, normalize_skill_names
from app.schemas.cv_profile import CVProfile, EducationEntry, ExperienceEntry
from app.services.cv_text import UnsupportedCVFormat, extract_raw_text

# -- CVProfile schema -------------------------------------------------------


def test_cv_profile_defaults_to_empty_lists() -> None:
    profile = CVProfile()

    assert profile.skills == []
    assert profile.experience == []
    assert profile.education == []
    assert profile.summary is None


def test_cv_profile_accepts_nested_experience_and_education() -> None:
    profile = CVProfile(
        current_title="Backend Engineer",
        skills=["Python", "SQL"],
        experience=[
            ExperienceEntry(
                title="Backend Engineer",
                company="Acme",
                start_date="Jan 2022",
                end_date="Present",
                is_current=True,
            )
        ],
        education=[
            EducationEntry(degree="B.Tech", institution="NIT", field_of_study="CS")
        ],
    )

    assert profile.experience[0].company == "Acme"
    assert profile.experience[0].is_current is True
    assert profile.education[0].degree == "B.Tech"


def test_cv_profile_serializes_nested_entries_as_plain_dicts() -> None:
    """This is exactly what app.services.cv_extraction stores into

    Profile.experience / Profile.education (JSONB columns) — a list of
    plain dicts, not Pydantic model instances, which JSONB cannot
    serialize on its own.
    """
    profile = CVProfile(
        experience=[ExperienceEntry(title="Engineer", company="Acme")]
    )

    dumped = [entry.model_dump(mode="json") for entry in profile.experience]

    assert dumped == [
        {
            "title": "Engineer",
            "company": "Acme",
            "start_date": None,
            "end_date": None,
            "start_year": None,
            "start_month": None,
            "end_year": None,
            "end_month": None,
            "is_current": False,
            "description": None,
        }
    ]


# -- skill normalization ------------------------------------------------


def test_normalize_lowercases_and_collapses_whitespace() -> None:
    assert normalize_skill_name("  Python   Programming  ") == "python programming"


def test_normalize_folds_known_punctuation_variants() -> None:
    assert normalize_skill_name("Node.js") == "nodejs"
    assert normalize_skill_name("C++") == "cpp"
    assert normalize_skill_name("C#") == "csharp"


def test_normalize_is_case_insensitive_for_aliases() -> None:
    assert normalize_skill_name("NODE.JS") == "nodejs"


def test_normalize_passes_through_names_with_no_alias() -> None:
    assert normalize_skill_name("Kubernetes") == "kubernetes"


# -- file text extraction ------------------------------------------------


def test_extract_docx_text_returns_paragraph_text() -> None:
    import io

    from docx import Document

    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("Backend Engineer with 5 years of experience.")

    buffer = io.BytesIO()
    document.save(buffer)

    text = extract_raw_text("docx", buffer.getvalue())

    assert "Jane Doe" in text
    assert "Backend Engineer with 5 years of experience." in text


def test_extract_pdf_text_is_empty_for_a_blank_page() -> None:
    """A well-formed PDF with a page but no text — the scanned-image case.

    Built with PdfWriter rather than a hand-written byte string. An
    earlier version of this test used the same minimal, structurally
    incomplete PDF bytes used elsewhere in this project purely to pass
    a magic-byte check (b"%PDF-..." with no real xref table). That is
    not the same thing as a scan: pypdf.PdfReader refuses to open it
    at all (PdfReadError: startxref not found), which extract_cv
    correctly treats as FAILED, not NO_TEXT_LAYER. A blank page is a
    file pypdf can genuinely open and find nothing to read in — the
    actual case this behavior needs to cover.
    """
    import io

    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)

    buffer = io.BytesIO()
    writer.write(buffer)

    assert extract_raw_text("pdf", buffer.getvalue()) == ""


def test_extract_pdf_text_raises_for_a_malformed_pdf() -> None:
    """A structurally broken file is distinct from a genuinely empty one.

    app.services.cv_extraction relies on this raising rather than
    returning "": a malformed PDF becomes FAILED (something is
    actually wrong with the file), while a blank or scanned page
    becomes NO_TEXT_LAYER (nothing is wrong, there was just nothing to
    read). Confirms the actual exception type pypdf raises for this
    shape of input, so a pypdf upgrade that silently changed it would
    be caught here rather than surfacing as a mysteriously
    uncategorized extraction failure later.
    """
    from pypdf.errors import PdfReadError

    malformed_pdf = (
        b"%PDF-1.4\n1 0 obj<</Type/Catalog>>endobj\ntrailer<</Root 1 0 R>>\n%%EOF"
    )

    with pytest.raises(PdfReadError):
        extract_raw_text("pdf", malformed_pdf)


def test_extract_raw_text_rejects_an_unknown_file_type() -> None:
    with pytest.raises(UnsupportedCVFormat):
        extract_raw_text("txt", b"plain text")


# -- normalize_skill_names ------------------------------------------------


def test_normalize_skill_names_folds_spelling_variants_together() -> None:
    """Proves the same skill spelled three different ways yields one key."""
    assert normalize_skill_names(["Node.js", "NodeJS", "node.js"]) == ["nodejs"]


def test_normalize_skill_names_preserves_first_seen_order() -> None:
    """Proves de-duplication doesn't discard the order a candidate lists skills in."""
    assert normalize_skill_names(["PyTorch", "FastAPI", "pytorch"]) == [
        "pytorch",
        "fastapi",
    ]


def test_normalize_skill_names_applies_punctuation_aliases() -> None:
    """Proves the per-name alias table is applied across the whole list."""
    assert normalize_skill_names(["C++", "CI/CD"]) == ["cpp", "cicd"]


def test_normalize_skill_names_drops_blank_entries() -> None:
    """Proves whitespace-only entries don't survive as an empty-string key."""
    assert normalize_skill_names(["Python", "", "   "]) == ["python"]
